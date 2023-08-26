# coding:utf-8

from docx import Document

doc = Document('文本.docx')  # doc -> docx

for p in doc.paragraphs:
    print(p.text)

for t in doc.tables:
    for row in t.rows:  # dir(t)
        _row_str = ''
        for cell in row.cells:
            _row_str += cell.text + ','
        print(_row_str)
