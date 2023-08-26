# coding:utf-8

from docx import Document
from docx.shared import Inches, RGBColor, Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT, WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE

doc = Document()

for i in doc.styles:
    if i.type == WD_STYLE_TYPE.TABLE:
        print(i.name)

style = doc.styles['Normal']
style.font.name = '微软雅黑'
# style.font.color.rgb = RGBColor(255, 255, 255)
style.font.size = Pt(16)


title = doc.add_heading('', 0)
title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
title.style.font.size = Pt(20)
_t = title.add_run('My title\n123')
_t.italic = True
_t.bold = True
_t.underline = True

p = doc.add_paragraph('欢迎来到这里学习python')
p.add_run('\n这是关于word生成的知识').italic = True
p.add_run('\npython很有意思')
p.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT

p1 = doc.add_paragraph()
p1.alignment = WD_ALIGN_PARAGRAPH.RIGHT
_p1 = p1.add_run()

image = _p1.add_picture('logo2020.png', width=Inches(2), height=Inches(2))

title = ['name', 'age', 'sex']

table = doc.add_table(rows=1, cols=3, style='Colorful Shading Accent 6')

title_cells = table.rows[0].cells
title_cells[0].text = title[0]
title_cells[1].text = title[1]
title_cells[2].text = title[2]

data = [
    ('xiaomu', '10', 'man'),
    ('dewei', '34', 'man'),
    ('xiaoman', '18', 'women')
]

for d in data:
    row_cells = table.add_row().cells
    print(d[0])
    row_cells[0].text = d[0]  # name
    row_cells[1].text = d[1]  # age
    row_cells[2].text = d[2]  # sex

doc.add_page_break()
title = doc.add_heading('My title2', 0)

doc.add_page_break()
title = doc.add_heading('My title3', 0)

doc.save('test.docx')
