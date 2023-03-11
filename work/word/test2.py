# coding:utf-8

import glob

from docx import Document

class ReadDoc(object):
    def __init__(self, path):
        self.doc = Document(path)
        self.p_text = ''
        self.table_text = ''

        self.get_para()
        self.get_table()

    def get_para(self):
        for p in self.doc.paragraphs:
            self.p_text += p.text + '\n'

    def get_table(self):
        for table in self.doc.tables:
            for row in table.rows:
                _cell_str = ''
                for cell in row.cells:
                    _cell_str += cell.text + ','
                self.table_text += _cell_str + '\n'


def search_word(path, targets):
    result = glob.glob(path)
    final_result = []

    for i in result:
        isuse = True
        if glob.os.path.isfile(i):
            if i.endswith('.docx'):
                doc = ReadDoc(i)
                p_text = doc.p_text
                t_text = doc.table_text
                all_text = p_text + t_text

                for target in targets:
                    if target not in all_text:
                        isuse = False
                        break

                if not isuse:
                    continue
                final_result.append(i)
    return final_result


if __name__ == '__main__':
    path = glob.os.path.join(glob.os.getcwd(), '*')
    res = search_word(path, ['python', 'golang', 'react', '最佳'])
    print(res)
